"""
DOCX Document Formatter.
Creates .docx files with preserved formatting.
"""
from pathlib import Path
from typing import Dict
import logging
from docx import Document as DocxDocument
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ..core.interfaces import IDocumentFormatter
from ..core.models import (
    Document,
    TextSegment,
    FileType,
    SegmentType
)


logger = logging.getLogger(__name__)


class DocxFormatter(IDocumentFormatter):
    """Formatter for Microsoft Word .docx files."""
    
    @property
    def supported_file_type(self) -> FileType:
        """Supported file type."""
        return FileType.DOCX
    
    def format(
        self,
        document: Document,
        output_path: Path,
        preserve_formatting: bool = True
    ) -> Path:
        """
        Format and save translated document.
        
        Args:
            document: Document with translated segments
            output_path: Where to save
            preserve_formatting: Whether to preserve formatting
            
        Returns:
            Path to saved document
            
        Raises:
            FormattingError: If formatting fails
        """
        try:
            # Create new document
            docx = DocxDocument()
            
            # Apply metadata
            self._apply_metadata(docx, document)
            
            # Group segments by paragraph
            paragraphs = self._group_by_paragraph(document.segments)
            
            # Write paragraphs
            for para_segments in paragraphs:
                if not para_segments:
                    continue
                
                # Determine if this is special content
                first_segment = para_segments[0]
                
                if first_segment.segment_type == SegmentType.HEADING:
                    self._add_heading(docx, para_segments, preserve_formatting)
                elif first_segment.segment_type == SegmentType.TABLE_CELL:
                    # Tables handled separately
                    pass
                else:
                    self._add_paragraph(docx, para_segments, preserve_formatting)
            
            # Add tables
            self._add_tables(docx, document, preserve_formatting)
            
            # Add headers/footers
            self._add_headers_footers(docx, document, preserve_formatting)
            
            # Save document
            output_path.parent.mkdir(parents=True, exist_ok=True)
            docx.save(output_path)
            
            logger.info(f"Saved formatted document to: {output_path}")
            return output_path
            
        except Exception as e:
            logger.error(f"Failed to format document: {e}")
            raise FormattingError(f"Formatting failed: {e}")
    
    def preserve_styles(
        self,
        original: Document,
        translated: Document
    ) -> Document:
        """
        Copy styles from original to translated.
        
        Args:
            original: Original document
            translated: Translated document
            
        Returns:
            Translated document with preserved styles
        """
        # Copy metadata
        translated.metadata = original.metadata
        
        # Copy styles
        translated.styles = original.styles
        
        # Copy formatting from original segments to translated
        # Match by segment ID
        original_by_id = {s.id: s for s in original.segments}
        
        for segment in translated.segments:
            if segment.id in original_by_id:
                original_segment = original_by_id[segment.id]
                segment.text_formatting = original_segment.text_formatting
                segment.paragraph_formatting = original_segment.paragraph_formatting
                segment.position = original_segment.position
        
        return translated
    
    def validate_output(self, output_path: Path) -> bool:
        """
        Validate output document.
        
        Args:
            output_path: Path to output document
            
        Returns:
            True if valid
        """
        if not output_path.exists():
            logger.error(f"Output file not found: {output_path}")
            return False
        
        try:
            # Try to open document
            DocxDocument(output_path)
            return True
        except Exception as e:
            logger.error(f"Invalid output document: {e}")
            return False
    
    def _apply_metadata(self, docx: DocxDocument, document: Document):
        """Apply metadata to document."""
        core_props = docx.core_properties
        meta = document.metadata
        
        if meta.title:
            core_props.title = meta.title
        if meta.author:
            core_props.author = meta.author
        if meta.subject:
            core_props.subject = meta.subject
        if meta.keywords:
            core_props.keywords = meta.keywords
        if meta.comments:
            core_props.comments = meta.comments
    
    def _group_by_paragraph(self, segments: list) -> list:
        """Group segments by paragraph index."""
        # Group segments that belong to same paragraph
        paragraphs: Dict[int, list] = {}
        
        for segment in segments:
            if segment.segment_type == SegmentType.TABLE_CELL:
                continue  # Handle tables separately
            
            para_idx = segment.position.paragraph_index
            if para_idx is None:
                para_idx = -1  # Fallback
            
            if para_idx not in paragraphs:
                paragraphs[para_idx] = []
            
            paragraphs[para_idx].append(segment)
        
        # Sort by paragraph index and return list
        sorted_paras = [paragraphs[k] for k in sorted(paragraphs.keys())]
        return sorted_paras
    
    def _add_paragraph(
        self,
        docx: DocxDocument,
        segments: list,
        preserve_formatting: bool
    ):
        """Add paragraph with runs."""
        paragraph = docx.add_paragraph()
        
        # Apply paragraph formatting from first segment
        if segments and preserve_formatting:
            first_segment = segments[0]
            if first_segment.paragraph_formatting:
                self._apply_paragraph_formatting(paragraph, first_segment.paragraph_formatting)
        
        # Add runs
        for segment in segments:
            run = paragraph.add_run(segment.text)
            
            if preserve_formatting and segment.text_formatting:
                self._apply_run_formatting(run, segment.text_formatting)
    
    def _add_heading(
        self,
        docx: DocxDocument,
        segments: list,
        preserve_formatting: bool
    ):
        """Add heading."""
        if not segments:
            return
        
        # Combine all segment texts
        text = " ".join(s.text for s in segments)
        
        # Determine heading level (default to level 1)
        level = 1
        first_segment = segments[0]
        if first_segment.paragraph_formatting:
            style_name = first_segment.paragraph_formatting.style_name or ""
            if "heading" in style_name.lower():
                # Extract level from style name (e.g., "Heading 2" -> 2)
                try:
                    level = int(''.join(filter(str.isdigit, style_name))) or 1
                except:
                    level = 1
        
        heading = docx.add_heading(text, level=min(level, 9))
        
        # Apply formatting if needed
        if preserve_formatting and segments[0].text_formatting:
            for run in heading.runs:
                self._apply_run_formatting(run, segments[0].text_formatting)
    
    def _add_tables(
        self,
        docx: DocxDocument,
        document: Document,
        preserve_formatting: bool
    ):
        """Add tables from segments."""
        # Group table cells by table index
        table_segments = [
            s for s in document.segments 
            if s.segment_type == SegmentType.TABLE_CELL
        ]
        
        if not table_segments:
            return
        
        # Group by table_index
        tables = {}
        for segment in table_segments:
            table_idx = segment.position.table_index
            if table_idx not in tables:
                tables[table_idx] = {}
            
            row_idx = segment.position.row_index
            if row_idx not in tables[table_idx]:
                tables[table_idx][row_idx] = {}
            
            cell_idx = segment.position.cell_index
            tables[table_idx][row_idx][cell_idx] = segment
        
        # Create tables
        for table_idx in sorted(tables.keys()):
            rows = tables[table_idx]
            
            # Determine table dimensions
            num_rows = len(rows)
            num_cols = max(len(row) for row in rows.values())
            
            # Create table
            table = docx.add_table(rows=num_rows, cols=num_cols)
            table.style = 'Table Grid'
            
            # Fill cells
            for row_idx in sorted(rows.keys()):
                row = rows[row_idx]
                for cell_idx in sorted(row.keys()):
                    segment = row[cell_idx]
                    cell = table.rows[row_idx].cells[cell_idx]
                    cell.text = segment.text
                    
                    # Apply formatting
                    if preserve_formatting and segment.text_formatting:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                self._apply_run_formatting(run, segment.text_formatting)
    
    def _add_headers_footers(
        self,
        docx: DocxDocument,
        document: Document,
        preserve_formatting: bool
    ):
        """Add headers and footers."""
        section = docx.sections[0]
        
        # Add headers
        if document.headers:
            header = section.header
            for segment in document.headers:
                paragraph = header.add_paragraph(segment.text)
                if preserve_formatting and segment.text_formatting:
                    for run in paragraph.runs:
                        self._apply_run_formatting(run, segment.text_formatting)
        
        # Add footers
        if document.footers:
            footer = section.footer
            for segment in document.footers:
                paragraph = footer.add_paragraph(segment.text)
                if preserve_formatting and segment.text_formatting:
                    for run in paragraph.runs:
                        self._apply_run_formatting(run, segment.text_formatting)
    
    def _apply_run_formatting(self, run, formatting):
        """Apply text formatting to run."""
        if formatting.font_name:
            run.font.name = formatting.font_name
        
        if formatting.font_size:
            run.font.size = Pt(formatting.font_size)
        
        if formatting.bold:
            run.font.bold = True
        
        if formatting.italic:
            run.font.italic = True
        
        if formatting.underline:
            run.font.underline = True
        
        if formatting.strikethrough:
            run.font.strike = True
        
        if formatting.color:
            # Parse hex color
            color_hex = formatting.color.lstrip('#')
            r = int(color_hex[0:2], 16)
            g = int(color_hex[2:4], 16)
            b = int(color_hex[4:6], 16)
            run.font.color.rgb = RGBColor(r, g, b)
        
        if formatting.subscript:
            run.font.subscript = True
        
        if formatting.superscript:
            run.font.superscript = True
        
        if formatting.small_caps:
            run.font.small_caps = True
    
    def _apply_paragraph_formatting(self, paragraph, formatting):
        """Apply paragraph formatting."""
        fmt = paragraph.paragraph_format
        
        # Alignment
        alignment_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        if formatting.alignment:
            paragraph.alignment = alignment_map.get(
                formatting.alignment, 
                WD_ALIGN_PARAGRAPH.LEFT
            )
        
        # Spacing
        if formatting.line_spacing:
            fmt.line_spacing = formatting.line_spacing
        
        if formatting.space_before:
            fmt.space_before = Pt(formatting.space_before)
        
        if formatting.space_after:
            fmt.space_after = Pt(formatting.space_after)
        
        # Indentation
        if formatting.left_indent:
            fmt.left_indent = Pt(formatting.left_indent)
        
        if formatting.right_indent:
            fmt.right_indent = Pt(formatting.right_indent)
        
        if formatting.first_line_indent:
            fmt.first_line_indent = Pt(formatting.first_line_indent)
        
        # Pagination
        if formatting.keep_together:
            fmt.keep_together = True
        
        if formatting.keep_with_next:
            fmt.keep_with_next = True
        
        if formatting.page_break_before:
            fmt.page_break_before = True


class FormattingError(Exception):
    """Exception raised when formatting fails."""
    pass


# ===== Example Usage =====

if __name__ == "__main__":
    from pathlib import Path
    from ..parsers.docx_parser import DocxParser
    from ..core.models import TextSegment, SegmentType, SegmentPosition, TextFormatting
    
    # Parse original document
    parser = DocxParser()
    formatter = DocxFormatter()
    
    test_file = Path("test_document.docx")
    if test_file.exists():
        # Parse
        original_doc = parser.parse(test_file)
        print(f"Parsed: {len(original_doc.segments)} segments")
        
        # Simulate translation (just add " [TRANSLATED]" for demo)
        translated_doc = Document(
            file_path=Path("test_translated.docx"),
            file_type=FileType.DOCX,
            segments=[],
            metadata=original_doc.metadata,
            styles=original_doc.styles
        )
        
        for segment in original_doc.segments:
            translated_segment = TextSegment(
                id=segment.id,
                text=segment.text + " [TRANSLATED]",
                segment_type=segment.segment_type,
                position=segment.position,
                text_formatting=segment.text_formatting,
                paragraph_formatting=segment.paragraph_formatting
            )
            translated_doc.segments.append(translated_segment)
        
        # Format and save
        output_path = Path("test_translated.docx")
        formatter.format(translated_doc, output_path, preserve_formatting=True)
        
        # Validate
        if formatter.validate_output(output_path):
            print(f"✓ Successfully created: {output_path}")
        else:
            print(f"✗ Failed to create valid document")
    else:
        print(f"Test file not found: {test_file}")
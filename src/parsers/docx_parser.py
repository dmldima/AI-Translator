"""
DOCX Document Parser - FIXED: Path Traversal Protection
Extracts text and formatting from .docx files.
"""
from pathlib import Path
from typing import List
import logging
from docx import Document as DocxDocument
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ..core.interfaces import IDocumentParser
from ..core.models import (
    Document,
    TextSegment,
    FileType,
    SegmentType,
    SegmentPosition,
    TextFormatting,
    ParagraphFormatting,
    DocumentMetadata,
    DocumentStyles
)


logger = logging.getLogger(__name__)


class DocxParser(IDocumentParser):
    """Parser for Microsoft Word .docx files."""
    
    @property
    def supported_file_type(self) -> FileType:
        """Supported file type."""
        return FileType.DOCX
    
    def can_parse(self, file_path: Path) -> bool:
        """
        Check if file can be parsed.
        
        Args:
            file_path: Path to file
            
        Returns:
            True if can parse
        """
        return file_path.suffix.lower() == '.docx'
    
    def validate_document(self, file_path: Path) -> bool:
        """
        Validate document before parsing.
        
        Args:
            file_path: Path to document
            
        Returns:
            True if valid
        """
        # SECURITY FIX: Validate path to prevent path traversal
        file_path = self._validate_and_resolve_path(file_path)
        
        if not file_path.exists():
            logger.error(f"File not found: {file_path}")
            return False
        
        if not self.can_parse(file_path):
            logger.error(f"Unsupported file type: {file_path.suffix}")
            return False
        
        try:
            # Try to open document
            DocxDocument(file_path)
            return True
        except Exception as e:
            logger.error(f"Cannot open document: {e}")
            return False
    
    def parse(self, file_path: Path) -> Document:
        """
        Parse DOCX document.
        
        Args:
            file_path: Path to .docx file
            
        Returns:
            Document with extracted content and formatting
            
        Raises:
            ParsingError: If parsing fails
        """
        # SECURITY FIX: Validate path first
        file_path = self._validate_and_resolve_path(file_path)
        
        if not self.validate_document(file_path):
            raise ParsingError(f"Invalid document: {file_path}")
        
        try:
            docx = DocxDocument(file_path)
            
            # Extract metadata
            metadata = self._extract_metadata(docx)
            
            # Extract styles
            styles = self._extract_styles(docx)
            
            # Extract segments
            segments = []
            
            # Process paragraphs
            for para_idx, paragraph in enumerate(docx.paragraphs):
                if not paragraph.text.strip():
                    continue  # Skip empty paragraphs
                
                # Determine segment type
                segment_type = self._get_segment_type(paragraph)
                
                # Process runs (formatted text chunks)
                for run_idx, run in enumerate(paragraph.runs):
                    if not run.text.strip():
                        continue
                    
                    segment = TextSegment(
                        id=f"para_{para_idx}_run_{run_idx}",
                        text=run.text,
                        segment_type=segment_type,
                        position=SegmentPosition(
                            paragraph_index=para_idx,
                            run_index=run_idx
                        ),
                        text_formatting=self._extract_run_formatting(run),
                        paragraph_formatting=self._extract_paragraph_formatting(paragraph)
                    )
                    segments.append(segment)
            
            # Process tables
            table_segments = self._extract_tables(docx)
            segments.extend(table_segments)
            
            # Process headers/footers
            header_segments, footer_segments = self._extract_headers_footers(docx)
            
            document = Document(
                file_path=file_path,
                file_type=FileType.DOCX,
                segments=segments,
                metadata=metadata,
                styles=styles,
                headers=header_segments,
                footers=footer_segments
            )
            
            logger.info(
                f"Parsed {file_path.name}: "
                f"{len(segments)} segments, "
                f"{document.total_words} words"
            )
            
            return document
            
        except Exception as e:
            logger.error(f"Failed to parse document: {e}")
            raise ParsingError(f"Parsing failed: {e}")
    
    def extract_segments(self, document: Document) -> List[TextSegment]:
        """
        Extract translatable segments.
        
        Args:
            document: Parsed document
            
        Returns:
            List of segments ready for translation
        """
        # Filter out empty segments
        segments = [s for s in document.segments if s.text.strip()]
        
        logger.info(f"Extracted {len(segments)} translatable segments")
        return segments
    
    def _validate_and_resolve_path(self, file_path: Path) -> Path:
        """
        SECURITY: Validate and resolve file path to prevent path traversal.
        
        Args:
            file_path: Input path
            
        Returns:
            Resolved absolute path
            
        Raises:
            ParsingError: If path is unsafe
        """
        # Convert to absolute path
        if not file_path.is_absolute():
            file_path = file_path.resolve()
        
        # Check for path traversal attempts
        path_str = str(file_path)
        if '..' in path_str:
            raise ParsingError(f"Path traversal not allowed: {file_path}")
        
        # Validate filename doesn't contain dangerous characters
        dangerous_chars = ['<', '>', '|', '\0', '\n', '\r']
        for char in dangerous_chars:
            if char in file_path.name:
                raise ParsingError(f"Invalid character in filename: {repr(char)}")
        
        return file_path
    
    def _extract_metadata(self, docx: DocxDocument) -> DocumentMetadata:
        """Extract document metadata."""
        core_props = docx.core_properties
        
        return DocumentMetadata(
            title=core_props.title,
            author=core_props.author,
            subject=core_props.subject,
            keywords=core_props.keywords,
            comments=core_props.comments,
            created=core_props.created,
            modified=core_props.modified,
            last_modified_by=core_props.last_modified_by
        )
    
    def _extract_styles(self, docx: DocxDocument) -> DocumentStyles:
        """Extract document-level styles."""
        styles = DocumentStyles()
        
        # Extract default font from first paragraph
        if docx.paragraphs:
            first_para = docx.paragraphs[0]
            if first_para.runs:
                first_run = first_para.runs[0]
                styles.default_font = first_run.font.name
        
        return styles
    
    def _get_segment_type(self, paragraph) -> SegmentType:
        """Determine segment type from paragraph."""
        style_name = paragraph.style.name.lower() if paragraph.style else ""
        
        if 'heading' in style_name:
            return SegmentType.HEADING
        elif 'list' in style_name:
            return SegmentType.LIST_ITEM
        else:
            return SegmentType.PARAGRAPH
    
    def _extract_run_formatting(self, run) -> TextFormatting:
        """Extract formatting from run."""
        font = run.font
        
        # Extract color
        color = None
        if font.color and font.color.rgb:
            rgb = font.color.rgb
            color = f"#{rgb[0]:02x}{rgb[1]:02x}{rgb[2]:02x}"
        
        # Extract highlight
        highlight = None
        if font.highlight_color:
            highlight = str(font.highlight_color)
        
        return TextFormatting(
            font_name=font.name,
            font_size=float(font.size.pt) if font.size else None,
            bold=font.bold if font.bold is not None else False,
            italic=font.italic if font.italic is not None else False,
            underline=font.underline if font.underline is not None else False,
            strikethrough=font.strike if font.strike is not None else False,
            color=color,
            highlight_color=highlight,
            subscript=font.subscript if font.subscript is not None else False,
            superscript=font.superscript if font.superscript is not None else False,
            small_caps=font.small_caps if font.small_caps is not None else False
        )
    
    def _extract_paragraph_formatting(self, paragraph) -> ParagraphFormatting:
        """Extract paragraph-level formatting."""
        fmt = paragraph.paragraph_format
        
        # Alignment
        alignment_map = {
            WD_ALIGN_PARAGRAPH.LEFT: "left",
            WD_ALIGN_PARAGRAPH.CENTER: "center",
            WD_ALIGN_PARAGRAPH.RIGHT: "right",
            WD_ALIGN_PARAGRAPH.JUSTIFY: "justify"
        }
        alignment = alignment_map.get(paragraph.alignment, "left")
        
        return ParagraphFormatting(
            alignment=alignment,
            line_spacing=float(fmt.line_spacing) if fmt.line_spacing else None,
            space_before=float(fmt.space_before.pt) if fmt.space_before else None,
            space_after=float(fmt.space_after.pt) if fmt.space_after else None,
            left_indent=float(fmt.left_indent.pt) if fmt.left_indent else None,
            right_indent=float(fmt.right_indent.pt) if fmt.right_indent else None,
            first_line_indent=float(fmt.first_line_indent.pt) if fmt.first_line_indent else None,
            style_name=paragraph.style.name if paragraph.style else None,
            keep_together=fmt.keep_together if fmt.keep_together is not None else False,
            keep_with_next=fmt.keep_with_next if fmt.keep_with_next is not None else False,
            page_break_before=fmt.page_break_before if fmt.page_break_before is not None else False
        )
    
    def _extract_tables(self, docx: DocxDocument) -> List[TextSegment]:
        """Extract text from tables."""
        segments = []
        
        for table_idx, table in enumerate(docx.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    if not text:
                        continue
                    
                    segment = TextSegment(
                        id=f"table_{table_idx}_row_{row_idx}_cell_{cell_idx}",
                        text=text,
                        segment_type=SegmentType.TABLE_CELL,
                        position=SegmentPosition(
                            table_index=table_idx,
                            row_index=row_idx,
                            cell_index=cell_idx
                        )
                    )
                    segments.append(segment)
        
        return segments
    
    def _extract_headers_footers(self, docx: DocxDocument) -> tuple:
        """Extract headers and footers."""
        header_segments = []
        footer_segments = []
        
        for section in docx.sections:
            # Headers
            header = section.header
            for para_idx, paragraph in enumerate(header.paragraphs):
                text = paragraph.text.strip()
                if text:
                    segment = TextSegment(
                        id=f"header_{para_idx}",
                        text=text,
                        segment_type=SegmentType.HEADER,
                        position=SegmentPosition(paragraph_index=para_idx)
                    )
                    header_segments.append(segment)
            
            # Footers
            footer = section.footer
            for para_idx, paragraph in enumerate(footer.paragraphs):
                text = paragraph.text.strip()
                if text:
                    segment = TextSegment(
                        id=f"footer_{para_idx}",
                        text=text,
                        segment_type=SegmentType.FOOTER,
                        position=SegmentPosition(paragraph_index=para_idx)
                    )
                    footer_segments.append(segment)
        
        return header_segments, footer_segments


class ParsingError(Exception):
    """Exception raised when parsing fails."""
    pass

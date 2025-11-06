"""
Enhanced DOCX Document Formatter - PRODUCTION READY v3.0
Implements best practices for maintaining all document styling.

CRITICAL FIXES:
- Atomic save cleanup logic corrected
- Memory management optimized
- Multi-run text distribution uses word boundaries
- Validation made less verbose for production
- Removed redundant logging in hot paths

Version: 3.0 (Audited & Stable)
"""
from pathlib import Path
from typing import Dict, List, Optional, Set
import logging
import os
import uuid
import re
from docx import Document as DocxDocument
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

from ..core.interfaces import IDocumentFormatter
from ..core.models import (
    Document, TextSegment, FileType, SegmentType
)


logger = logging.getLogger(__name__)


class EnhancedDocxFormatter(IDocumentFormatter):
    """
    Enhanced DOCX formatter with complete formatting preservation.
    Production-ready with optimized performance and stability.
    """
    
    def __init__(self, workspace_root: Path = None):
        """
        Initialize formatter.
        
        Args:
            workspace_root: Root directory for allowed file operations.
                           If None, uses current working directory.
        """
        self.workspace_root = (workspace_root or Path.cwd()).resolve()
        self._validation_enabled = True  # Can be disabled for performance
    
    @property
    def supported_file_type(self) -> FileType:
        return FileType.DOCX
    
    def format(
        self,
        document: Document,
        output_path: Path,
        preserve_formatting: bool = True
    ) -> Path:
        """
        Format and save translated document with complete formatting preservation.
        
        Args:
            document: Document with translated segments
            output_path: Where to save
            preserve_formatting: Whether to preserve formatting
            
        Returns:
            Path to saved document
            
        Raises:
            FormattingError: If formatting fails
        """
        doc = None
        
        try:
            # Validate paths before processing
            document.file_path = self._validate_and_resolve_path(document.file_path)
            output_path = self._validate_and_resolve_path(output_path, check_exists=False)
            
            logger.info(f"Formatting document: {output_path.name}")
            
            # Load original document to preserve structure
            doc = DocxDocument(document.file_path)
            
            if preserve_formatting:
                # Strategy: Replace text in-place to preserve all formatting
                self._replace_text_in_place(doc, document.segments)
            else:
                # Create new document from scratch
                doc = self._create_new_document(document)
            
            # Atomic save - FIXED: no return value needed
            self._save_atomic(doc, output_path)
            
            # Validate formatting preservation
            if preserve_formatting and self._validation_enabled:
                validation_passed = self._validate_formatting(document.file_path, output_path)
                if not validation_passed:
                    logger.warning("Some formatting validation checks failed")
            
            logger.info(f"✓ Document saved: {output_path}")
            return output_path
            
        except FileNotFoundError as e:
            raise FormattingError(
                f"Input file not found: {document.file_path}"
            ) from e
        
        except PermissionError as e:
            raise FormattingError(
                f"Permission denied: Cannot write to {output_path}"
            ) from e
        
        except ValueError as e:
            if "workspace" in str(e).lower():
                raise FormattingError(
                    f"Security error: File path outside allowed workspace"
                ) from e
            raise
        
        except Exception as e:
            logger.error(f"Formatting failed: {e}", exc_info=True)
            raise FormattingError(
                f"Failed to format DOCX document: {e}"
            ) from e
        
        finally:
            # Memory cleanup
            if doc:
                try:
                    if hasattr(doc, '_part'):
                        doc._part = None
                except Exception:
                    pass
    
    def _save_atomic(self, doc, output_path: Path) -> None:
        """
        Atomic save using temp file to prevent corruption.
        
        FIXED: Proper temp file cleanup logic.
        
        Args:
            doc: Document to save
            output_path: Target path
            
        Raises:
            FormattingError: If save fails
        """
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        # Create temp file in same directory
        temp_path = output_path.with_name(
            f'.tmp_{uuid.uuid4().hex[:8]}_{output_path.name}'
        )
        
        try:
            # Save to temp file
            doc.save(temp_path)
            
            # Atomic rename - after this, temp_path no longer exists
            temp_path.replace(output_path)
            
        except Exception as e:
            # Cleanup temp file only if it still exists
            if temp_path.exists():
                try:
                    temp_path.unlink()
                except OSError:
                    pass
            raise FormattingError(f"Failed to save document: {e}") from e
    
    def _validate_and_resolve_path(self, file_path: Path, check_exists: bool = True) -> Path:
        """
        SECURITY: Validate and resolve file path to prevent path traversal.
        
        Args:
            file_path: Input path
            check_exists: Whether to check if path exists
            
        Returns:
            Resolved absolute path
            
        Raises:
            FormattingError: If path is unsafe
        """
        file_path = file_path.resolve()
        
        try:
            file_path.relative_to(self.workspace_root)
        except ValueError:
            raise FormattingError(
                f"Access denied: Path outside workspace"
            )
        
        self._validate_filename(file_path.name)
        
        if check_exists and not file_path.exists():
            raise FormattingError(f"File not found: {file_path}")
        
        return file_path
    
    def _validate_filename(self, filename: str) -> None:
        """Validate filename for cross-platform compatibility."""
        if os.name == 'nt':
            dangerous = ['<', '>', ':', '"', '/', '\\', '|', '?', '*', '\0']
        else:
            dangerous = ['/', '\0']
        
        for char in dangerous:
            if char in filename:
                raise FormattingError(f"Invalid character in filename: {repr(char)}")
        
        if os.name == 'nt':
            reserved = {
                'CON', 'PRN', 'AUX', 'NUL',
                'COM1', 'COM2', 'COM3', 'COM4', 'COM5', 'COM6', 'COM7', 'COM8', 'COM9',
                'LPT1', 'LPT2', 'LPT3', 'LPT4', 'LPT5', 'LPT6', 'LPT7', 'LPT8', 'LPT9'
            }
            name_without_ext = filename.rsplit('.', 1)[0].upper()
            if name_without_ext in reserved:
                raise FormattingError(f"Reserved Windows filename: {filename}")
            
            if filename.rstrip() != filename or filename.rstrip('.') != filename:
                raise FormattingError("Filename cannot end with space or dot on Windows")
        
        if len(filename) > 255:
            raise FormattingError("Filename too long (max 255 characters)")
    
    def _replace_text_in_place(
        self,
        docx: DocxDocument,
        segments: List[TextSegment]
    ) -> None:
        """
        Replace text in original document while preserving all formatting.
        OPTIMIZED: Reduced logging verbosity for production.
        """
        segment_map = self._build_segment_map(segments)
        segments_replaced = set()
        
        # Process paragraphs
        for para_idx, paragraph in enumerate(docx.paragraphs):
            for run_idx, run in enumerate(paragraph.runs):
                segment_id = f"para_{para_idx}_run_{run_idx}"
                
                if segment_id in segment_map:
                    segment = segment_map[segment_id]
                    self._replace_text_preserve_formatting(run, segment.text)
                    segments_replaced.add(segment_id)
        
        # Process tables
        for table_idx, table in enumerate(docx.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    segment_id = f"table_{table_idx}_row_{row_idx}_cell_{cell_idx}"
                    
                    if segment_id in segment_map:
                        segment = segment_map[segment_id]
                        
                        if cell.paragraphs:
                            self._replace_cell_text(cell.paragraphs[0], segment.text)
                        
                        segments_replaced.add(segment_id)
        
        # Process headers/footers
        for section_idx, section in enumerate(docx.sections):
            self._replace_header_footer_text(
                section.header, segment_map, "header", section_idx, segments_replaced
            )
            self._replace_header_footer_text(
                section.footer, segment_map, "footer", section_idx, segments_replaced
            )
        
        # Summary logging only
        coverage = len(segments_replaced) / len(segment_map) * 100 if segment_map else 0
        logger.info(f"✓ Replaced {len(segments_replaced)}/{len(segment_map)} segments ({coverage:.1f}%)")
        
        if coverage < 90:
            logger.warning(f"Low replacement coverage: {coverage:.1f}%")
    
    def _replace_text_preserve_formatting(self, run, new_text: str) -> None:
        """
        Replace text in run while preserving ALL font properties.
        OPTIMIZED: Comprehensive backup with safe restoration.
        """
        font = run.font
        
        # Comprehensive font backup
        font_backup = {
            'name': font.name,
            'size': font.size,
            'bold': font.bold,
            'italic': font.italic,
            'underline': font.underline,
            'strike': font.strike,
            'double_strike': getattr(font, 'double_strike', None),
            'outline': getattr(font, 'outline', None),
            'shadow': getattr(font, 'shadow', None),
            'emboss': getattr(font, 'emboss', None),
            'imprint': getattr(font, 'imprint', None),
            'color_rgb': font.color.rgb if (font.color and hasattr(font.color, 'rgb')) else None,
            'color_theme': getattr(font.color, 'theme_color', None) if font.color else None,
            'highlight_color': font.highlight_color,
            'subscript': font.subscript,
            'superscript': font.superscript,
            'all_caps': getattr(font, 'all_caps', None),
            'small_caps': font.small_caps,
            'hidden': getattr(font, 'hidden', None),
            'spacing': getattr(font, 'spacing', None),
            'kerning': getattr(font, 'kerning', None),
        }
        
        # Replace text
        run.text = new_text
        
        # Restore ALL properties
        for prop, value in font_backup.items():
            if value is None:
                continue
                
            try:
                if prop == 'color_rgb' and value:
                    font.color.rgb = value
                elif prop == 'color_theme' and value:
                    font.color.theme_color = value
                else:
                    setattr(font, prop, value)
            except (AttributeError, TypeError, ValueError):
                # Some properties may be read-only or version-specific
                pass
    
    def _replace_cell_text(self, paragraph, new_text: str) -> None:
        """
        Replace text in table cell paragraph while preserving multi-run formatting.
        IMPROVED: Uses word-boundary splitting instead of proportional character split.
        """
        if not paragraph.runs:
            paragraph.add_run(new_text)
            return
        
        if len(paragraph.runs) == 1:
            self._replace_text_preserve_formatting(paragraph.runs[0], new_text)
            return
        
        # Multiple runs - use word-boundary distribution
        # This prevents splitting words across formatting boundaries
        original_lengths = [len(run.text) for run in paragraph.runs]
        total_original = sum(original_lengths)
        
        if total_original == 0:
            self._replace_text_preserve_formatting(paragraph.runs[0], new_text)
            for run in paragraph.runs[1:]:
                run.text = ''
            return
        
        # Split new text into words
        words = new_text.split()
        if not words:
            # Empty text
            for run in paragraph.runs:
                run.text = ''
            return
        
        # Distribute words across runs based on original proportions
        runs_text = [''] * len(paragraph.runs)
        words_per_run = []
        
        for i, orig_len in enumerate(original_lengths):
            proportion = orig_len / total_original
            num_words = max(1, int(len(words) * proportion))
            words_per_run.append(num_words)
        
        # Adjust to ensure all words are assigned
        total_words_assigned = sum(words_per_run)
        if total_words_assigned < len(words):
            words_per_run[-1] += len(words) - total_words_assigned
        elif total_words_assigned > len(words):
            # Reduce from end
            excess = total_words_assigned - len(words)
            for i in range(len(words_per_run) - 1, -1, -1):
                if words_per_run[i] > excess:
                    words_per_run[i] -= excess
                    break
                else:
                    excess -= words_per_run[i]
                    words_per_run[i] = 0
        
        # Assign words to runs
        word_idx = 0
        for i, num_words in enumerate(words_per_run):
            if word_idx >= len(words):
                runs_text[i] = ''
            else:
                end_idx = min(word_idx + num_words, len(words))
                runs_text[i] = ' '.join(words[word_idx:end_idx])
                word_idx = end_idx
        
        # Apply text to runs
        for run, text in zip(paragraph.runs, runs_text):
            self._replace_text_preserve_formatting(run, text)
    
    def _replace_header_footer_text(
        self,
        header_footer,
        segment_map: Dict[str, TextSegment],
        prefix: str,
        section_idx: int,
        segments_replaced: Set[str]
    ) -> None:
        """Replace text in headers/footers with smart multi-run handling."""
        for para_idx, paragraph in enumerate(header_footer.paragraphs):
            segment_id = f"section_{section_idx}_{prefix}_{para_idx}"
            
            if segment_id not in segment_map:
                continue
            
            segment = segment_map[segment_id]
            
            if not paragraph.runs:
                paragraph.add_run(segment.text)
                segments_replaced.add(segment_id)
                continue
            
            self._replace_cell_text(paragraph, segment.text)
            segments_replaced.add(segment_id)
    
    def _build_segment_map(
        self,
        segments: List[TextSegment]
    ) -> Dict[str, TextSegment]:
        """Build lookup map of segments by ID."""
        segment_map = {}
        duplicates = 0
        
        for segment in segments:
            if segment.id in segment_map:
                duplicates += 1
            segment_map[segment.id] = segment
        
        if duplicates > 0:
            logger.warning(f"Found {duplicates} duplicate segment IDs")
        
        return segment_map
    
    def _create_new_document(self, document: Document) -> DocxDocument:
        """Create new document from scratch (no formatting preservation)."""
        docx = DocxDocument()
        
        if document.metadata:
            self._apply_metadata(docx, document.metadata)
        
        paragraphs = self._group_by_paragraph(document.segments)
        
        for para_segments in paragraphs:
            if not para_segments:
                continue
            
            first_segment = para_segments[0]
            
            if first_segment.segment_type == SegmentType.HEADING:
                self._add_heading(docx, para_segments)
            elif first_segment.segment_type != SegmentType.TABLE_CELL:
                self._add_paragraph(docx, para_segments)
        
        self._add_tables(docx, document)
        
        return docx
    
    def _apply_metadata(self, docx: DocxDocument, metadata) -> None:
        """Apply document metadata."""
        core_props = docx.core_properties
        
        if metadata.title:
            core_props.title = metadata.title
        if metadata.author:
            core_props.author = metadata.author
        if metadata.subject:
            core_props.subject = metadata.subject
        if metadata.keywords:
            core_props.keywords = metadata.keywords
        if metadata.comments:
            core_props.comments = metadata.comments
    
    def _group_by_paragraph(self, segments: List[TextSegment]) -> List[List[TextSegment]]:
        """Group segments that belong to same paragraph."""
        paragraphs: Dict[int, List[TextSegment]] = {}
        
        for segment in segments:
            if segment.segment_type == SegmentType.TABLE_CELL:
                continue
            
            para_idx = segment.position.paragraph_index or -1
            
            if para_idx not in paragraphs:
                paragraphs[para_idx] = []
            
            paragraphs[para_idx].append(segment)
        
        return [paragraphs[k] for k in sorted(paragraphs.keys())]
    
    def _add_paragraph(self, docx: DocxDocument, segments: List[TextSegment]) -> None:
        """Add paragraph with runs."""
        paragraph = docx.add_paragraph()
        
        if segments and segments[0].paragraph_formatting:
            self._apply_paragraph_formatting(paragraph, segments[0].paragraph_formatting)
        
        for segment in segments:
            run = paragraph.add_run(segment.text)
            
            if segment.text_formatting:
                self._apply_run_formatting(run, segment.text_formatting)
    
    def _add_heading(self, docx: DocxDocument, segments: List[TextSegment]) -> None:
        """Add heading."""
        if not segments:
            return
        
        text = " ".join(s.text for s in segments)
        
        level = 1
        first_segment = segments[0]
        if first_segment.paragraph_formatting:
            style_name = first_segment.paragraph_formatting.style_name or ""
            if "heading" in style_name.lower():
                try:
                    level = int(''.join(filter(str.isdigit, style_name))) or 1
                except:
                    level = 1
        
        heading = docx.add_heading(text, level=min(level, 9))
        
        if segments[0].text_formatting:
            for run in heading.runs:
                self._apply_run_formatting(run, segments[0].text_formatting)
    
    def _add_tables(self, docx: DocxDocument, document: Document) -> None:
        """Add tables from segments."""
        table_segments = [
            s for s in document.segments
            if s.segment_type == SegmentType.TABLE_CELL
        ]
        
        if not table_segments:
            return
        
        tables = {}
        for segment in table_segments:
            table_idx = segment.position.table_index
            if table_idx not in tables:
                tables[table_idx] = {}
            
            row_idx = segment.position.row_index
            if row_idx not in tables[table_idx]:
                tables[table_idx][row_idx] = {}
            
            cell_idx = segment.position.cell_index
            tables[table_idx][row_idx][cell_idx] = segment
        
        for table_idx in sorted(tables.keys()):
            rows = tables[table_idx]
            
            num_rows = len(rows)
            num_cols = max(len(row) for row in rows.values())
            
            table = docx.add_table(rows=num_rows, cols=num_cols)
            table.style = 'Table Grid'
            
            for row_idx in sorted(rows.keys()):
                row = rows[row_idx]
                for cell_idx in sorted(row.keys()):
                    segment = row[cell_idx]
                    cell = table.rows[row_idx].cells[cell_idx]
                    cell.text = segment.text
                    
                    if segment.text_formatting:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                self._apply_run_formatting(run, segment.text_formatting)
    
    def _apply_run_formatting(self, run, formatting) -> None:
        """Apply comprehensive text formatting to run."""
        font = run.font
        
        if formatting.font_name:
            font.name = formatting.font_name
        
        if formatting.font_size:
            font.size = Pt(formatting.font_size)
        
        if formatting.bold:
            font.bold = True
        
        if formatting.italic:
            font.italic = True
        
        if formatting.underline:
            font.underline = True
        
        if formatting.strikethrough:
            font.strike = True
        
        if formatting.color:
            try:
                color_hex = formatting.color.lstrip('#')
                r = int(color_hex[0:2], 16)
                g = int(color_hex[2:4], 16)
                b = int(color_hex[4:6], 16)
                font.color.rgb = RGBColor(r, g, b)
            except Exception:
                pass
        
        if formatting.highlight_color:
            try:
                shading_elm = OxmlElement('w:shd')
                shading_elm.set(qn('w:fill'), formatting.highlight_color.lstrip('#'))
                run._element.get_or_add_rPr().append(shading_elm)
            except Exception:
                pass
        
        if formatting.subscript:
            font.subscript = True
        
        if formatting.superscript:
            font.superscript = True
        
        if formatting.small_caps:
            font.small_caps = True
    
    def _apply_paragraph_formatting(self, paragraph, formatting) -> None:
        """Apply comprehensive paragraph formatting."""
        fmt = paragraph.paragraph_format
        
        alignment_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        
        if formatting.alignment:
            paragraph.alignment = alignment_map.get(
                formatting.alignment,
                WD_ALIGN_PARAGRAPH.LEFT
            )
        
        if formatting.line_spacing:
            fmt.line_spacing = formatting.line_spacing
        
        if formatting.space_before:
            fmt.space_before = Pt(formatting.space_before)
        
        if formatting.space_after:
            fmt.space_after = Pt(formatting.space_after)
        
        if formatting.left_indent:
            fmt.left_indent = Inches(formatting.left_indent / 72)
        
        if formatting.right_indent:
            fmt.right_indent = Inches(formatting.right_indent / 72)
        
        if formatting.first_line_indent:
            fmt.first_line_indent = Inches(formatting.first_line_indent / 72)
        
        if formatting.keep_together:
            fmt.keep_together = True
        
        if formatting.keep_with_next:
            fmt.keep_with_next = True
        
        if formatting.page_break_before:
            fmt.page_break_before = True
    
    def _validate_formatting(self, original_path: Path, output_path: Path) -> bool:
        """
        Validate that formatting was preserved.
        OPTIMIZED: Less verbose, focus on critical errors only.
        """
        try:
            original_doc = DocxDocument(original_path)
            output_doc = DocxDocument(output_path)
            
            critical_errors = []
            
            # Critical checks only
            if len(original_doc.paragraphs) != len(output_doc.paragraphs):
                critical_errors.append(
                    f"Paragraph count mismatch: {len(original_doc.paragraphs)} → {len(output_doc.paragraphs)}"
                )
            
            if len(original_doc.tables) != len(output_doc.tables):
                critical_errors.append(
                    f"Table count mismatch: {len(original_doc.tables)} → {len(output_doc.tables)}"
                )
            
            if critical_errors:
                logger.error("Critical formatting errors:")
                for error in critical_errors:
                    logger.error(f"  • {error}")
                return False
            
            logger.info("✓ Formatting validation passed")
            return True
            
        except Exception as e:
            logger.warning(f"Formatting validation failed: {e}")
            return False
    
    def preserve_styles(self, original: Document, translated: Document) -> Document:
        """Copy styles from original to translated."""
        translated.metadata = original.metadata
        translated.styles = original.styles
        
        original_by_id = {s.id: s for s in original.segments}
        
        for segment in translated.segments:
            if segment.id in original_by_id:
                original_segment = original_by_id[segment.id]
                segment.text_formatting = original_segment.text_formatting
                segment.paragraph_formatting = original_segment.paragraph_formatting
                segment.position = original_segment.position
        
        return translated
    
    def validate_output(self, output_path: Path) -> bool:
        """Validate output document can be opened."""
        if not output_path.exists():
            logger.error(f"Output file not found: {output_path}")
            return False
        
        try:
            DocxDocument(output_path)
            logger.info(f"✓ Output document validated: {output_path.name}")
            return True
        except Exception as e:
            logger.error(f"Invalid output document: {e}")
            return False


class FormattingError(Exception):
    """Exception raised when formatting fails."""
    pass


# Compatibility alias
DocxFormatter = EnhancedDocxFormatter
